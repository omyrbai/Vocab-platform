from collections.abc import Sequence

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt

from app.schemas.term import TermRead


class FlashcardService:

    COLUMN_WIDTH = Inches(4.13)
    ROW_HEIGHT = Inches(2.90)

    def create_docx(
            self,
            terms: Sequence[TermRead],
            output_path: str,
    ) -> None:
        """
        Create a printable flashcard Word document.
        """

        document = Document()

        section = document.sections[0]

        # A4
        section.page_width = Mm(210)
        section.page_height = Mm(297)

        # Zero margins
        section.top_margin = Mm(0)
        section.bottom_margin = Mm(0)
        section.left_margin = Mm(0)
        section.right_margin = Mm(0)

        section.header_distance = Mm(0)
        section.footer_distance = Mm(0)

        # Process 8 flashcards at a time.
        for page_start in range(0, len(terms), 8):

            page_terms = terms[
                page_start:page_start + 8
            ]

            # if page_start > 0:
            #     document.add_page_break()

            # FRONT
            self._add_front_page(
                document,
                page_terms,
            )

            # BACK
            # document.add_page_break()

            self._add_back_page(
                document,
                page_terms,
            )

        document.save(output_path)

    def _add_front_page(
            self,
            document: Document,
            terms: Sequence[TermRead],
    ) -> None:

        table = self._create_table(document)

        for index, term in enumerate(terms):

            row = index // 2
            col = index % 2

            cell = table.cell(row, col)

            self._configure_cell(cell)

            paragraph = cell.paragraphs[0]

            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER
            )

            self._remove_paragraph_spacing(
                paragraph
            )

            run = paragraph.add_run(
                term.term
            )

            run.font.name = "Candara"
            run.font.size = Pt(28)
            run.bold = True

        self._set_table_borders(table)

    def _add_back_page(
            self,
            document: Document,
            terms: Sequence[TermRead],
    ) -> None:

        table = self._create_table(document)

        for index, term in enumerate(terms):

            # IMPORTANT:
            #
            # Front:
            #
            # 1 | 2
            # 3 | 4
            #
            # Back must be mirrored:
            #
            # 2 | 1
            # 4 | 3
            #
            # Therefore:
            row = index // 2
            col = 1 - (index % 2)

            cell = table.cell(row, col)

            self._configure_cell(
                cell,
                back_page=True,
            )

            self._add_term_information(
                cell,
                term,
            )

        self._set_table_borders(table)

    def _create_table(
            self,
            document: Document,
    ):
        """
        Create an exactly 4 × 2 flashcard table.
        """

        table = document.add_table(
            rows=4,
            cols=2,
        )

        table.autofit = False

        # Remove table indentation.
        tblPr = table._tbl.tblPr

        tblInd = tblPr.first_child_found_in(
            "w:tblInd"
        )

        if tblInd is not None:
            tblPr.remove(tblInd)

        # Set table width.
        table.width = Inches(8.26)

        for row in table.rows:

            row.height = self.ROW_HEIGHT
            row.height_rule = 1  # exact

            for cell in row.cells:

                cell.width = (
                    self.COLUMN_WIDTH
                )

        return table

    def _configure_cell(
            self,
            cell,
            back_page: bool = False,
    ) -> None:

        cell.width = self.COLUMN_WIDTH
        cell.height = self.ROW_HEIGHT

        cell.vertical_alignment = (
            WD_CELL_VERTICAL_ALIGNMENT.CENTER
        )

        # Remove cell internal margins.
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()

        tcMar = tcPr.first_child_found_in(
            "w:tcMar"
        )

        if tcMar is None:
            tcMar = OxmlElement(
                "w:tcMar"
            )
            tcPr.append(tcMar)

        for margin in (
                "top",
                "left",
                "bottom",
                "right",
        ):
            node = tcMar.find(
                qn(f"w:{margin}")
            )

            if node is None:
                node = OxmlElement(
                    f"w:{margin}"
                )
                tcMar.append(node)

            if back_page and margin == "left":
                value = "400"
            else:
                value = "0"

            node.set(
                qn("w:w"),
                value,
            )

            node.set(
                qn("w:type"),
                "dxa",
            )

    def _add_term_information(
            self,
            cell,
            term: TermRead,
    ) -> None:

        paragraph = cell.paragraphs[0]

        paragraph.alignment = (
            WD_ALIGN_PARAGRAPH.LEFT
        )

        self._remove_paragraph_spacing(
            paragraph
        )

        if term.pronunciation:

            run = paragraph.add_run(
                f"{term.pronunciation}\n"
            )

            run.font.name = "Candara"
            run.font.size = Pt(14)

        run = paragraph.add_run(
            f"{term.definition}\n"
        )

        run.font.name = "Candara"
        run.font.size = Pt(14)

        run = paragraph.add_run(
            f"{term.example}\n"
        )

        run.font.name = "Candara"
        run.font.size = Pt(14)

        run = paragraph.add_run(
            term.translation
        )

        run.font.name = "Candara"
        run.font.size = Pt(14)

    def _remove_paragraph_spacing(
            self,
            paragraph,
    ) -> None:

        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1

    def _set_table_borders(
            self,
            table,
    ) -> None:

        tbl = table._tbl
        tblPr = tbl.tblPr

        borders = tblPr.first_child_found_in(
            "w:tblBorders"
        )

        if borders is None:

            borders = OxmlElement(
                "w:tblBorders"
            )

            tblPr.append(borders)

        for edge in (
            "top",
            "left",
            "bottom",
            "right",
            "insideH",
            "insideV",
        ):

            element = borders.find(
                qn(f"w:{edge}")
            )

            if element is None:

                element = OxmlElement(
                    f"w:{edge}"
                )

                borders.append(element)

            element.set(
                qn("w:val"),
                "single",
            )

            element.set(
                qn("w:sz"),
                "8",
            )

            element.set(
                qn("w:space"),
                "0",
            )